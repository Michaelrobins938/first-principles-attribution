import os
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, Frame, NextPageTemplate
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus.flowables import PageBreak


def read_markdown_file(file_path):
    """Read the markdown file and return its content."""
    with open(file_path, 'r', encoding='utf-8') as file:
        content = file.read()
    return content


def clean_text_for_pdf(text):
    """Clean text to remove or replace problematic Unicode characters."""
    replacements = {
        'φ': '(phi)',  # Greek phi
        '∑': 'sum',    # Summation symbol
        '∈': 'in',     # Element of
        '∀': 'for all', # For all
        '→': '->',     # Right arrow
        '≠': '!=',     # Not equal
        '≤': '<=',     # Less than or equal
        '≥': '>=',     # Greater than or equal
        '×': '*',      # Multiplication
        '∞': 'inf',    # Infinity
        'α': '(alpha)', # Greek alpha
        'β': '(beta)',  # Greek beta
        'γ': '(gamma)', # Greek gamma
        'δ': '(delta)', # Greek delta
        'ε': '(epsilon)', # Greek epsilon
        'θ': '(theta)', # Greek theta
        'λ': '(lambda)', # Greek lambda
        'μ': '(mu)',    # Greek mu
        'π': '(pi)',    # Greek pi
        'σ': '(sigma)', # Greek sigma
        'τ': '(tau)',   # Greek tau
        'Φ': '(Phi)',   # Greek Phi
        'Δ': 'Delta',   # Greek Delta
        'Γ': 'Gamma',   # Greek Gamma
        'Λ': 'Lambda',  # Greek Lambda
        'Ω': 'Omega',   # Greek Omega
        '√': 'sqrt',    # Square root
        '≈': '~=',      # Approximately equal
        '∝': 'prop',    # Proportional to
        '∇': 'del',     # Nabla/del operator
        '∂': 'partial', # Partial derivative
        '∫': 'int',     # Integral
        '∬': 'int2',    # Double integral
        '∭': 'int3',    # Triple integral
        '∮': 'contour', # Contour integral
    }
    
    for old, new in replacements.items():
        text = text.replace(old, new)
    
    return text


def parse_markdown_to_structure(content):
    """Parse markdown content into structured data."""
    lines = content.split('\n')
    structure = []
    current_section = None
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if line.startswith('---'):
            # Skip frontmatter
            i += 1
            while i < len(lines) and not lines[i].startswith('---'):
                i += 1
            i += 1
            continue
        
        if line.startswith('# '):
            # Main title
            title = line[2:].strip()
            structure.append({'type': 'title', 'content': title})
        elif line.startswith('## '):
            # Main heading
            heading = line[3:].strip()
            structure.append({'type': 'heading1', 'content': heading})
        elif line.startswith('### '):
            # Sub-heading
            heading = line[4:].strip()
            structure.append({'type': 'heading2', 'content': heading})
        elif line.startswith('**') and line.endswith('**') and ':' in line:
            # Bold key-value pairs
            structure.append({'type': 'bold', 'content': line.strip('* ')})
        elif line.startswith('|') and '|' in line:
            # Table
            table_data = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [cell.strip() for cell in lines[i].strip('|').split('|')]
                table_data.append(row)
                i += 1
            structure.append({'type': 'table', 'content': table_data})
            i -= 1  # Adjust for the loop increment
        elif line.startswith('```'):
            # Code block
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            structure.append({'type': 'code', 'content': '\n'.join(code_lines)})
        elif line == '':
            # Empty line
            structure.append({'type': 'empty'})
        else:
            # Regular paragraph
            paragraph = line
            # Look ahead for continuation
            j = i + 1
            while j < len(lines) and lines[j].strip() != '' and not lines[j].strip().startswith('#') and not lines[j].strip().startswith('|') and not lines[j].strip().startswith('```'):
                paragraph += ' ' + lines[j].strip()
                j += 1
            
            structure.append({'type': 'paragraph', 'content': paragraph})
            i = j - 1  # Adjust index
        
        i += 1
    
    return structure


class WhitePaperDOCX:
    def __init__(self):
        self.doc = Document()
        self.setup_styles()
    
    def setup_styles(self):
        """Setup custom styles for the document."""
        # Title style
        title_style = self.doc.styles['Title']
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(16)
        title_style.font.bold = True
        
        # Heading 1 style
        h1_style = self.doc.styles['Heading 1']
        h1_style.font.name = 'Arial'
        h1_style.font.size = Pt(14)
        h1_style.font.bold = True
        h1_style.paragraph_format.space_after = Pt(12)
        
        # Heading 2 style
        h2_style = self.doc.styles['Heading 2']
        h2_style.font.name = 'Arial'
        h2_style.font.size = Pt(12)
        h2_style.font.bold = True
        h2_style.paragraph_format.space_after = Pt(10)
        
        # Normal paragraph style
        normal_style = self.doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(11)
        normal_style.paragraph_format.line_spacing = 1.5
        normal_style.paragraph_format.space_after = Pt(6)
    
    def add_title_page(self):
        """Add a professional title page."""
        # Add title
        title_para = self.doc.add_paragraph()
        title_run = title_para.add_run('A First-Principles Hybrid Attribution Framework\nIntegrating Probabilistic Path Modeling, Cooperative Game Theory, and Psychographic Transition Priors')
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(16)
        title_run.font.bold = True
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Empty line
        
        # Add subtitle
        subtitle_para = self.doc.add_paragraph()
        subtitle_run = subtitle_para.add_run('Technical Whitepaper v2.0.0')
        subtitle_run.font.name = 'Arial'
        subtitle_run.font.size = Pt(14)
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_paragraph()  # Empty line
        
        # Add metadata
        meta_para = self.doc.add_paragraph()
        meta_run = meta_para.add_run(
            'Classification: Methodology Specification / Decision Science\n'
            'Status: Production-Ready / Frozen\n'
            'Document Version: 2.0.0\n'
            'Last Updated: January 23, 2026\n'
            'Organization: Advanced Analytics Division\n'
            'Contact: attribution-engine@organization.com'
        )
        meta_run.font.name = 'Arial'
        meta_run.font.size = Pt(11)
        meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()
    
    def add_content_from_structure(self, structure):
        """Add content to the document based on the parsed structure."""
        for item in structure:
            if item['type'] == 'title':
                para = self.doc.add_paragraph(item['content'], style='Title')
            elif item['type'] == 'heading1':
                para = self.doc.add_paragraph(item['content'], style='Heading 1')
            elif item['type'] == 'heading2':
                para = self.doc.add_paragraph(item['content'], style='Heading 2')
            elif item['type'] == 'paragraph':
                para = self.doc.add_paragraph(item['content'], style='Normal')
            elif item['type'] == 'bold':
                para = self.doc.add_paragraph(style='Normal')
                run = para.add_run(item['content'])
                run.bold = True
            elif item['type'] == 'table':
                # Add table to document
                if len(item['content']) > 1:  # Ensure there's a header
                    table = self.doc.add_table(rows=len(item['content']), cols=len(item['content'][0]))
                    table.style = 'Table Grid'
                    
                    for i, row_data in enumerate(item['content']):
                        for j, cell_data in enumerate(row_data):
                            cell = table.cell(i, j)
                            cell.text = cell_data
                            
                            # Format header row differently
                            if i == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
            elif item['type'] == 'code':
                para = self.doc.add_paragraph(style='Normal')
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = para.add_run(item['content'])
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif item['type'] == 'empty':
                # Add spacing equivalent to empty line
                self.doc.add_paragraph('')
    
    def save(self, filename):
        """Save the document."""
        self.doc.save(filename)


class WhitePaperPDF:
    def __init__(self, filename):
        self.filename = filename
        self.doc = SimpleDocTemplate(filename, pagesize=A4, 
                                   topMargin=72, bottomMargin=72, 
                                   leftMargin=72, rightMargin=72)
        self.styles = getSampleStyleSheet()
        self.elements = []
        
        # Create custom styles
        self.title_style = ParagraphStyle(
            'CustomTitle',
            parent=self.styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=TA_CENTER,  # Center alignment
            textColor=colors.black,
            fontName='Helvetica-Bold'
        )
        
        self.subtitle_style = ParagraphStyle(
            'Subtitle',
            parent=self.styles['Normal'],
            fontSize=14,
            spaceAfter=20,
            alignment=TA_CENTER,  # Center alignment
            textColor=colors.black,
            fontName='Helvetica-Bold'
        )
        
        self.metadata_style = ParagraphStyle(
            'Metadata',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceAfter=20,
            alignment=TA_CENTER,  # Center alignment
            textColor=colors.black,
            fontName='Helvetica'
        )
        
        self.heading1_style = ParagraphStyle(
            'CustomHeading1',
            parent=self.styles['Heading1'],
            fontSize=14,
            spaceBefore=12,
            spaceAfter=6,
            borderWidth=1,
            borderColor=colors.black,
            borderPadding=5,
            fontName='Helvetica-Bold'
        )
        
        self.heading2_style = ParagraphStyle(
            'CustomHeading2',
            parent=self.styles['Heading2'],
            fontSize=12,
            spaceBefore=10,
            spaceAfter=4,
            fontName='Helvetica-Bold'
        )
        
        self.normal_style = ParagraphStyle(
            'CustomNormal',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=6,
            alignment=TA_JUSTIFY,
            fontName='Helvetica'
        )
        
        self.bold_style = ParagraphStyle(
            'BoldStyle',
            parent=self.styles['Normal'],
            fontSize=11,
            spaceBefore=6,
            spaceAfter=6,
            alignment=TA_LEFT,
            fontName='Helvetica-Bold'
        )
        
        self.code_style = ParagraphStyle(
            'CodeStyle',
            parent=self.styles['Normal'],
            fontName='Courier',
            fontSize=10,
            spaceBefore=6,
            spaceAfter=6,
            leftIndent=20,
            backColor=colors.lightgrey
        )
    
    def add_title_page(self):
        """Add a professional title page."""
        # Title
        title = Paragraph('A First-Principles Hybrid Attribution Framework', self.title_style)
        self.elements.append(title)
        
        subtitle = Paragraph('Integrating Probabilistic Path Modeling, Cooperative Game Theory, and Psychographic Transition Priors', self.title_style)
        self.elements.append(subtitle)
        
        # Subtitle
        subtitle_text = Paragraph('Technical Whitepaper v2.0.0', self.subtitle_style)
        self.elements.append(subtitle_text)
        
        # Metadata
        meta_text = Paragraph(
            'Classification: Methodology Specification / Decision Science<br/>' +
            'Status: Production-Ready / Frozen<br/>' +
            'Document Version: 2.0.0<br/>' +
            'Last Updated: January 23, 2026<br/>' +
            'Organization: Advanced Analytics Division<br/>' +
            'Contact: attribution-engine@organization.com',
            self.metadata_style
        )
        self.elements.append(meta_text)
        
        # Add a horizontal line
        self.elements.append(Spacer(1, 20))
        
        # Add page break
        self.elements.append(PageBreak())
    
    def add_content_from_structure(self, structure):
        """Add content to the PDF based on the parsed structure."""
        for item in structure:
            if item['type'] == 'title':
                cleaned_content = clean_text_for_pdf(item['content'])
                p = Paragraph(cleaned_content, self.title_style)
                self.elements.append(p)
            elif item['type'] == 'heading1':
                cleaned_content = clean_text_for_pdf(item['content'])
                p = Paragraph(cleaned_content, self.heading1_style)
                self.elements.append(p)
            elif item['type'] == 'heading2':
                cleaned_content = clean_text_for_pdf(item['content'])
                p = Paragraph(cleaned_content, self.heading2_style)
                self.elements.append(p)
            elif item['type'] == 'paragraph':
                cleaned_content = clean_text_for_pdf(item['content'])
                # Replace newlines with <br/> for ReportLab
                cleaned_content = cleaned_content.replace('\n', '<br/>')
                p = Paragraph(cleaned_content, self.normal_style)
                self.elements.append(p)
            elif item['type'] == 'bold':
                cleaned_content = clean_text_for_pdf(item['content'])
                p = Paragraph(f'<b>{cleaned_content}</b>', self.bold_style)
                self.elements.append(p)
            elif item['type'] == 'table':
                # Create table
                if len(item['content']) > 0:
                    # Clean table data
                    cleaned_table_data = []
                    for row in item['content']:
                        cleaned_row = [clean_text_for_pdf(str(cell)) for cell in row]
                        cleaned_table_data.append(cleaned_row)
                    
                    table = Table(cleaned_table_data)
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#444444")),  # Dark gray header
                        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, 0), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                        ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                        ('FONTNAME', (0, 1), (-1, -1), 'Helvetica'),
                        ('FONTSIZE', (0, 1), (-1, -1), 8),
                        ('GRID', (0, 0), (-1, -1), 1, colors.black),
                        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                    ]))
                    self.elements.append(table)
                    self.elements.append(Spacer(1, 12))  # Add space after table
            elif item['type'] == 'code':
                cleaned_content = clean_text_for_pdf(item['content'])
                p = Paragraph(f'<font face="Courier">{cleaned_content}</font>', self.code_style)
                self.elements.append(p)
            elif item['type'] == 'empty':
                self.elements.append(Spacer(1, 12))
    
    def save(self):
        """Build and save the PDF."""
        self.doc.build(self.elements)


def main():
    # Define file paths
    md_file_path = r"D:\first-principles-attribution-repo\docs\FIRST_PRINCIPLES_ATTRIBUTION_WHITEPAPER_PROFESSIONAL_FORMAT.md"
    docx_output_path = r"D:\first-principles-attribution-repo\docs\FIRST_PRINCIPLES_ATTRIBUTION_WHITEPAPER_PROFESSIONAL_FORMAT.docx"
    pdf_output_path = r"D:\first-principles-attribution-repo\docs\FIRST_PRINCIPLES_ATTRIBUTION_WHITEPAPER_PROFESSIONAL_FORMAT_POLISHED.pdf"
    
    # Read the markdown file
    content = read_markdown_file(md_file_path)
    
    # Parse the markdown content
    structure = parse_markdown_to_structure(content)
    
    # Create and save PDF with improved formatting
    print("Creating polished PDF file...")
    pdf_writer = WhitePaperPDF(pdf_output_path)
    pdf_writer.add_title_page()
    pdf_writer.add_content_from_structure(structure)
    pdf_writer.save()
    print(f"Polished PDF file saved to: {pdf_output_path}")
    
    print("\nPolished PDF creation completed successfully!")


if __name__ == "__main__":
    main()